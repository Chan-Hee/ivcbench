#!/usr/bin/env python
"""Build the submission-ready Supp_Results.docx from supp/results_supplementary.md.

Styled to match Supp_TableS1.docx: US-Letter, 1in margins, Times New Roman; title 14pt bold,
'Supplementary Note S#' headings 12pt bold, body 10pt with bold lead labels. Notes are renumbered
S1..S5 (md) -> S2..S6 (docx) to continue past the existing curation 'Supp Note S1'.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Emu

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "supp" / "results_supplementary.md"
OUT = ROOT.parent / "Supp_Results.docx"
MS_TITLE = "Toward Immune Virtual Cells: Benchmarking Perturbation Prediction Models for Immunology"
INTRO = ("This document provides supporting detail for the benchmark Results section of the main text, "
         "including the per-unit detail of the two added models, CellOT (neural optimal transport) and "
         "chemCPA (native chemistry-aware CPA). Each note expands a corresponding main-text subsection. "
         "All values are generated under leak-audited splits and are summarized in main "
         "Tables 3 to 6 and Figures 2 to 6. The CellOT donor-transfer result revises the donor-axis verdict "
         "to a model-family-dependent outcome (as stated in the main text); chemCPA and CellOT on Kang "
         "corroborate the existing verdicts.")

doc = Document()
sec = doc.sections[0]
sec.page_width = Emu(7772400); sec.page_height = Emu(10058400)
sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Emu(914400)
doc.styles['Normal'].font.name = 'Times New Roman'; doc.styles['Normal'].font.size = Pt(10)


def title(text, pt):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True
    r.font.name = 'Times New Roman'; r.font.size = Pt(pt); return p

def plain(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(10); return p

def body_labelled(text):
    p = doc.add_paragraph()
    m = re.match(r"^(.{1,80}?\.)\s", text)
    if m:
        r1 = p.add_run(m.group(1)); r1.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(10)
        r2 = p.add_run(" " + text[m.end():]); r2.font.name = 'Times New Roman'; r2.font.size = Pt(10)
    else:
        r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return p


title("Supplementary Materials for the Immune-Aware Perturbation Benchmark", 14)
plain(MS_TITLE)
plain(INTRO)

note_re = re.compile(r"^##\s+S(\d+)\.\s+(.*)$")
for ln in (l.rstrip("\n") for l in MD.open()):
    m = note_re.match(ln)
    if m:
        title(f"Supplementary Note S{m.group(1)}. {m.group(2)}", 12)
    elif ln.startswith("# ") or not ln.strip() or ln.startswith("This file holds") or ln.startswith("This document provides"):
        continue
    else:
        body_labelled(ln.strip())

doc.save(OUT)
d = Document(OUT)
heads = [p.text for p in d.paragraphs if p.text.startswith("Supplementary Note S")]
em = sum(p.text.count('—') for p in d.paragraphs)
print(f"wrote {OUT.name}: {len(d.paragraphs)} paras, {len(heads)} notes, em-dash={em}")
for h in heads:
    print("  ", h)
