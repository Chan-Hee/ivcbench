#!/usr/bin/env python
"""Apply Soskic donor-axis result edits to results_section.md and the manuscript docx.

Run after:
  1. scripts/soskic_donor_postprocess.py
  2. figure_framework.py, figure_ranking.py, figure_donor_decision.py
  3. refresh_docx_figures.py and update_docx_captions.py (optional; this script also updates captions)
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "results" / "_paper" / "results_section.md"
DOCX = ROOT.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx"
SUMMARY = ROOT / "results" / "C2" / "soskic_donor_bootstrap_summary.csv"
FOREST = ROOT / "figures" / "source_data" / "soskic_donor_forest.csv"


def fmt(v: float, nd: int = 3) -> str:
    return f"{v:+.{nd}f}".replace("-", "−")


def vals() -> dict:
    s = pd.read_csv(SUMMARY)
    f = pd.read_csv(FOREST)
    pear = s[s.metric == "pearson_delta"].iloc[0]
    out = dict(
        n=int(pear.n),
        mean=float(pear["mean"]),
        lo=float(pear.lo),
        hi=float(pear.hi),
        pct=float(pear.pct_positive),
        sc=float(f.scGen.mean()),
        primary=float(f.primary_baseline.mean()),
    )
    if out["lo"] > 0:
        out["verdict"] = "conditioning exceeds baseline"
        out["branch"] = "positive"
    elif out["hi"] <= 0:
        out["verdict"] = "baseline not exceeded"
        out["branch"] = "negative"
    else:
        out["verdict"] = "competitive / interval crosses zero"
        out["branch"] = "mixed"
    return out


def prose(v: dict) -> dict[str, str]:
    delta = f"{fmt(v['mean'])} [{fmt(v['lo'])}, {fmt(v['hi'])}]"
    pct = f"{100 * v['pct']:.1f}%"
    if v["branch"] == "positive":
        donor_read = (
            f"On the dedicated Soskic donor-resolved substrate, scGen exceeded the pre-specified primary "
            f"baseline under leave-one-donor-out (mean gap {delta}; {pct} donors positive; n = {v['n']}). "
            "This positive donor result conflicts with any headline stating that conditioning helps only "
            "on cell-context transfer; those headline edits should be proposed explicitly rather than "
            "silently folded into the manuscript."
        )
        overview = (
            "Across this benchmark panel, the advantage of perturbation conditioning was not a property of "
            "a model family alone; it depended on the generalization axis being tested (Figure 3, Table 6). "
            "Conditioning reached or exceeded the simple baselines for seen-perturbation cell-context "
            "transfer and, in the newly added Soskic analysis, for donor-resolved CD4 activation LODO; it "
            "did not exceed a non-conditioned training-mean or donor shift when the task was to extrapolate "
            "to an unseen perturbation or to stress modality. Reading the method × split landscape by "
            "column (axis) rather than by row (model) is therefore essential: because the same model family "
            "changes verdict across columns, a leaderboard pooled over splits would average these opposing "
            "signs into a single misleading rank. This is a results-level statement restricted to the "
            "evaluated immune datasets, splits, model implementations, and baselines; the axis-level "
            "endpoints below are treated as primary readouts, and finer per-lineage, per-marker, and "
            "per-perturbation breakdowns as exploratory."
        )
        synthesis = (
            "Across the evaluated axes, perturbation conditioning exceeded simple baselines for "
            "seen-perturbation cell-context transfer, most clearly in OP3 and locally in Kang, and for "
            "Soskic donor-resolved CD4 activation LODO. It did not exceed simple baselines for unseen-gene, "
            "unseen-compound, or modality generalization in these splits. Program ranking and response "
            "magnitude also separated, showing that immune-aware metrics change the interpretation of "
            "apparent model performance. These findings justify axis-stratified reporting rather than "
            "pooled leaderboards."
        )
    else:
        donor_read = (
            f"On the dedicated Soskic donor-resolved substrate, conditioning did not exceed the "
            f"pre-specified primary baseline under leave-one-donor-out: scGen averaged {v['sc']:.3f}, "
            f"the per-donor primary baseline averaged {v['primary']:.3f}, and the donor-bootstrap gap was "
            f"{delta} with {pct} donors positive (n = {v['n']}; Figure 6b, Table 6). Kang is therefore "
            "retained as the random-split optimism control rather than the main donor-axis evidence. In "
            "Kang, random cell splits inflated Pearson-Delta by +0.017 at the observation level (paired "
            "Wilcoxon p = 2.2 × 10⁻³) and by a donor-unit cluster-bootstrap interval of [−0.0003, +0.033] "
            "over 8 donors, so random splits remain a small but suggestive optimism control."
        )
        overview = (
            "Across this benchmark panel, the advantage of perturbation conditioning was not a property of "
            "a model family alone; it depended on the generalization axis being tested (Figure 3, Table 6). "
            "Conditioning reached or exceeded the simple baselines when a seen perturbation had to be "
            "transferred into an unseen immune cell context, but did not exceed a non-conditioned "
            "training-mean or donor shift when the task was to extrapolate to an unseen perturbation, to "
            "stress modality, or to generalize to an unseen donor in Soskic. Reading the method × split "
            "landscape by column (axis) rather than by row (model) is therefore essential: because the same "
            "model family changes verdict across columns, a leaderboard pooled over splits would average "
            "these opposing signs into a single misleading rank. This is a results-level statement "
            "restricted to the evaluated immune datasets, splits, model implementations, and baselines; "
            "the axis-level endpoints below are treated as primary readouts, and finer per-lineage, "
            "per-marker, and per-perturbation breakdowns as exploratory."
        )
        synthesis = (
            "Across the evaluated axes, perturbation conditioning exceeded simple baselines only for "
            "seen-perturbation cell-context transfer, most clearly in OP3 and locally in Kang. It did not "
            "exceed simple baselines for unseen-gene, unseen-compound, or modality generalization in these "
            "splits, nor for donor generalization in the Soskic leave-one-donor-out analysis. Program "
            "ranking and response magnitude also separated, showing that immune-aware metrics change the "
            "interpretation of apparent model performance. These findings justify axis-stratified reporting "
            "rather than pooled leaderboards."
        )
    return {
        "overview": overview,
        "donor": donor_read,
        "synthesis": synthesis,
        "bridge": (
            "These five task clusters map onto the four generalization axes through five benchmark "
            "components (Table 3). Cytokine-response prediction (Kang) supplies the cell-context split and "
            "a random-split donor control; donor-aware activation dynamics is now evaluated directly on "
            "Soskic CD4 activation with leave-one-donor-out; gene-intervention prediction (the "
            "primary-human-T CRISPR collection) supplies the perturbation axis; complex-context perturbation "
            "(Frangieh) supplies the modality axis; and small-molecule perturbation (OP3) supplies a second "
            "cell-context split together with the unseen-compound perturbation split. The previous Soskic "
            "exclusion reason does not hold for the processed 0h-versus-16h data: paired control-perturbed "
            "donor structure and matched simple baselines are present. Every axis therefore has at least "
            "one evaluated dataset, and no benchmark component is left unmapped."
        ),
        "panel": (
            "The evaluated panel comprises five benchmark components, each mapped to the axis it most "
            "directly stresses: the Kang IFN-β PBMC dataset for cytokine-driven cell-context transfer and a "
            "random-split donor control; the Soskic CD4 activation atlas for donor-resolved 0h-to-16h "
            "leave-one-donor-out; the OP3 PBMC small-molecule atlas for chemical cell-context transfer and "
            "unseen-compound extrapolation; a primary-human-T CRISPR collection (five datasets) for "
            "unseen-gene extrapolation; and Frangieh Perturb-CITE-seq for an RNA-versus-CITE modality "
            "stress test in a melanoma–tumour-infiltrating-lymphocyte (TIL) immune-evasion co-culture. "
            "This panel does not exhaust the curated dataset landscape described above; it operationalizes "
            "the evaluation principles on the resources for which the required structure and baselines were "
            "available. Where the figures and table use compact labels, they denote these same datasets "
            "(Kang, Soskic, OP3, primary-T CRISPR, Frangieh)."
        ),
        "metric": (
            "The four immune-aware metrics are not all defined on every component (Table 4). "
            "Response-direction fidelity (Pearson-Delta) and distributional fidelity (E-distance) are "
            "computed on all five datasets. Immune-program fidelity (AUCell-Delta) is computed wherever a "
            "curated transcriptional immune gene set is defined: the type-I-interferon set for Kang, "
            "T-cell activation / IL2-STAT5 / type-I and type-II IFN modules for Soskic, three response "
            "programs (ISG, NF-kB, effector) for OP3, and five T-cell programs for the primary-T CRISPR "
            "collection, but not for Frangieh, whose held readout is the CITE-seq surface-protein panel "
            "rather than a transcriptional program. Split-stratified generalization robustness applies to "
            "every endpoint by construction. The Figure 2 schematic therefore depicts the full metric menu, "
            "while Table 4 records where each metric is actually instantiated."
        ),
        "limitations": (
            "Several limitations bound these conclusions. First, statistical power is modest: the "
            "resampling unit is the biological replicate — 6 immune lineages, 5 datasets, 106 Soskic "
            "donors, or 8 Kang donors for the random-split control — so the axis-level endpoints are "
            "consistent-direction effects rather than high-powered tests. Second, the evaluated panel is a "
            "deliberate funnel from the surveyed landscape (Tables 3–5): datasets were admitted only where "
            "paired control-perturbed structure and matched baselines existed; Soskic now satisfies this "
            "criterion for the 0h-versus-16h donor axis, whereas temporal saturation and several adjacent "
            "resources remain outside the present benchmark. Third, model coverage is partial — eight "
            "conditioned models of the roughly thirty surveyed were executed (Table 5), so the absence of a "
            "published method here is not evidence against it. Fourth, the single positive axis in the "
            "negative-donor reading (seen-compound cell-context transfer on OP3) is also the comparatively "
            "easier task — interpolating an already-seen perturbation across lineages — so it should not be "
            "read as general evidence that conditioning transfers."
        ),
    }


def table6_cells(v: dict) -> list[str]:
    return [
        "Donor",
        "Soskic CD4 activation, leave-one-donor-out (0h->16h)",
        f"donor (n = {v['n']})",
        f"{v['primary']:.3f} (best of training-mean/donor shift)",
        f"scGen {v['sc']:.3f}",
        "E-distance and AUCell-Delta in source data",
        f"{fmt(v['mean'])} [{fmt(v['lo'])}, {fmt(v['hi'])}]",
        f"cluster bootstrap over donors; {100 * v['pct']:.1f}% positive",
        v["verdict"],
    ]


def update_md(v: dict, p: dict[str, str]) -> None:
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    out = []
    t6 = "| " + " | ".join(table6_cells(v)) + " |"
    for line in lines:
        if line.startswith("These five task clusters map onto"):
            out.append(p["bridge"])
        elif line.startswith("| Donor-aware activation dynamics |"):
            out.append("| Donor-aware activation dynamics | Soskic CD4 activation (106 paired donors) | Donor (leave-one-donor-out) |")
        elif line.startswith("The evaluated panel comprises four benchmark components"):
            out.append(p["panel"])
        elif line.startswith("The four immune-aware metrics are not all defined"):
            out.append(p["metric"])
        elif line.startswith("| Kang IFN-β PBMC | ✓ | ✓ |"):
            out.append(line)
            out.append("| Soskic CD4 activation | ✓ | ✓ | ✓ (T-cell activation; IL2/STAT5; type-I/II IFN) | ✓ |")
        elif line.startswith("| scGen | latent |"):
            out.append("| scGen | latent | Kang, Soskic, Frangieh; CRISPR, OP3 | native (Kang, Soskic, Frangieh) / adapted (CRISPR, OP3) |")
        elif line.startswith("Across this benchmark panel, the advantage"):
            out.append(p["overview"])
        elif line.startswith("Conditioning did not exceed the simple baseline on the donor axis"):
            out.append(p["donor"])
        elif line.startswith("Several limitations bound these conclusions"):
            out.append(p["limitations"])
        elif line.startswith("Across the evaluated axes, perturbation conditioning exceeded"):
            out.append(p["synthesis"])
        elif line.startswith("| Donor | Kang, leave-one-donor-out"):
            out.append(t6)
        elif line.startswith("**Figure 2."):
            out.append("**Figure 2. Benchmark framework.** The benchmark evaluates immune perturbation prediction along four generalization axes, scores predictions with up to four immune-aware metrics (applied per dataset as defined; Table 4), and is instantiated on a panel of five public datasets. **(a)** The four generalization axes that define what each task holds out: cell-context (unseen cell type), perturbation (unseen gene or compound), modality (RNA versus surface protein), and donor (unseen individual). **(b)** The four immune-aware metric axes, each applied where it is defined for a dataset (Table 4): response-direction fidelity (Pearson-Delta), distributional fidelity (E-distance), immune-program fidelity (AUCell-Delta on curated immune gene sets), and split-stratified generalization robustness. **(c)** The evaluated dataset panel and the axis each most directly stresses: Kang IFN-beta PBMC, Soskic CD4 activation, the OP3 PBMC small-molecule atlas, the primary-human-T CRISPR collection, and Frangieh Perturb-CITE-seq. Schematic; no error bars.")
        elif line.startswith("**Figure 3."):
            out.append("**Figure 3. Method × split performance landscape.** Heatmap of Pearson-Delta for every applicable (model, split) pair; rows are models grouped by family (simple baselines, then conditioned families), columns are evaluation splits grouped by the generalization axis they test, and cell fill encodes the metric value on a diverging scale (rust < 0 < blue, centred at the control/no-difference level). **(a)** The cell-context, modality, and donor axes — splits are leave-one-lineage-out (cell-context), unseen-knockout RNA versus CITE (modality), and leave-one-donor-out (donor; Kang and Soskic columns); conditioned families visibly separate from the simple baselines only under the cell-context columns. **(b)** The perturbation axis — leave-one-gene-out and unseen-compound splits — where no conditioned family separates from the simple baselines. Corner triangles flag adapted (re-fit) and shared-simple-baseline cells; pale-grey cells are not applicable. Datasets, as tagged above each column block: Kang IFN-beta PBMC; Soskic CD4 activation; the primary-human-T CRISPR collection (Chen, McCutcheon, Schmidt, Shifrut); Frangieh Perturb-CITE-seq (RNA and surface-protein readouts); and the OP3 compound atlas. The statistical decisions are in Table 6.")
        elif line.startswith("**Figure 6."):
            out.append(f"**Figure 6. Donor generalization and axis-level decisions (Kang, Soskic).** **(a)** Kang random-split optimism control: per-baseline inflation (random-split minus leak-proof leave-one-donor-out Pearson-Delta) across 8 Kang donors, with each baseline's mean and the donor-unit bootstrap summary. **(b)** Soskic CD4 activation leave-one-donor-out is the primary donor-axis evidence: each point is one held-out donor (n = {v['n']}), showing scGen minus the pre-specified primary baseline on response-gene Pearson-Delta; the vertical summary line and band report the donor-bootstrap mean {fmt(v['mean'])} and 95% CI [{fmt(v['lo'])}, {fmt(v['hi'])}]. **(c)** Axis-level summary of conditioned-minus-primary-baseline decisions, with uncertainty shown where the biological-unit bootstrap is the prespecified readout. Kang random-split inflation and Soskic donor generalization are separate donor-axis questions; Table 6 gives the statistical decisions.")
        else:
            out.append(line)
    MD.write_text("\n".join(out) + "\n", encoding="utf-8")


def set_para_text(para, text: str) -> None:
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    para.add_run(text)


def update_docx(v: dict, p: dict[str, str]) -> None:
    if not DOCX.exists():
        raise SystemExit(f"docx not found: {DOCX}")
    backup = DOCX.with_name(DOCX.stem + ".pre-soskic.docx")
    if not backup.exists():
        shutil.copy2(DOCX, backup)
    doc = Document(str(DOCX))
    replacements = [
        ("These five task clusters map onto", p["bridge"]),
        ("The evaluated panel comprises four benchmark components", p["panel"]),
        ("The four immune-aware metrics are not all defined", p["metric"]),
        ("Across this benchmark panel, the advantage", p["overview"]),
        ("Conditioning did not exceed the simple baseline on the donor axis", p["donor"]),
        ("Several limitations bound these conclusions", p["limitations"]),
        ("Across the evaluated axes, perturbation conditioning exceeded", p["synthesis"]),
    ]
    for para in doc.paragraphs:
        txt = para.text
        for prefix, repl in replacements:
            if txt.startswith(prefix):
                set_para_text(para, repl)
                break

    # Data availability before References.
    if not any(p_.text.strip() == "Data availability" for p_ in doc.paragraphs):
        for para in doc.paragraphs:
            if para.text.strip() == "References":
                para.insert_paragraph_before(
                    "The Soskic 2022 processed CD4 activation h5ad files used here were obtained from "
                    "the Trynka lab public object store; the donor-axis benchmark uses the paired 0h "
                    "resting and 16h highly active CD4 files with 106 processed paired donors.")
                para.insert_paragraph_before("Data availability")
                break

    # Tables 1 and 3-6.
    t1 = doc.tables[0]
    for row in t1.rows:
        if any("Soskic" in c.text for c in row.cells):
            row.cells[2].text = "CD4+ T cells (sorted; processed paired-donor benchmark subset)"
            row.cells[4].text = "Anti-CD3/CD28 Dynabeads; benchmark contrast 0 h resting -> 16 h highly active CD4"
            row.cells[5].text = "scRNA-seq (10x Chromium)"
            row.cells[6].text = (
                "Benchmark subset: 106 paired donors for 0 h resting to 16 h highly active CD4 contrast "
                "(processed HVG h5ad files)"
            )

    t3, t4, t5, t6 = doc.tables[2], doc.tables[3], doc.tables[4], doc.tables[5]
    for row in t3.rows:
        if row.cells[0].text.strip().startswith("Donor-aware activation dynamics"):
            row.cells[1].text = "Soskic CD4 activation (106 paired donors)"
            row.cells[2].text = "Donor (leave-one-donor-out)"
    if not any("Soskic CD4 activation" in r.cells[0].text for r in t4.rows):
        r = t4.add_row().cells
        vals4 = ["Soskic CD4 activation", "✓", "✓",
                 "✓ (T-cell activation; IL2/STAT5; type-I/II IFN)", "✓"]
        for c, x in zip(r, vals4):
            c.text = x
    for row in t5.rows:
        if row.cells[0].text.strip() == "scGen":
            row.cells[2].text = "Kang, Soskic, Frangieh; CRISPR, OP3"
            row.cells[3].text = "native (Kang, Soskic, Frangieh) / adapted (CRISPR, OP3)"
    cells6 = table6_cells(v)
    for row in t6.rows:
        if row.cells[0].text.strip() == "Donor" and row.cells[1].text.strip().startswith("Kang"):
            for c, x in zip(row.cells, cells6):
                c.text = x

    doc.save(str(DOCX))
    print(f"updated {DOCX}")
    print(f"backup {backup}")


def main() -> None:
    if not SUMMARY.exists() or not FOREST.exists():
        raise SystemExit("Run scripts/soskic_donor_postprocess.py before applying manuscript edits.")
    v = vals()
    p = prose(v)
    update_md(v, p)
    update_docx(v, p)
    if v["branch"] == "positive":
        (ROOT / "reports" / "soskic_headline_rewrite_needed.md").write_text(
            "Soskic donor LODO is positive. Revise any 'only cell-context' headline/abstract wording by "
            "explicit tracked-diff proposal, not silent automatic replacement.\n",
            encoding="utf-8",
        )
    print(f"Soskic branch: {v['branch']} ({fmt(v['mean'])} [{fmt(v['lo'])}, {fmt(v['hi'])}])")


if __name__ == "__main__":
    main()
