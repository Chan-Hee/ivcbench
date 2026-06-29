#!/usr/bin/env python3
"""Refresh the eight benchmark figures embedded in a results .docx IN PLACE.

The .docx layout (headings, body prose, Table 3, footnotes, figure legends) was hand-approved; this script
must NOT touch any of it. It only swaps the eight benchmark figure images for their freshly-rendered
versions. Each figure is identified by the "Figure N." caption that follows it (N = 1..8, the v2 plate with
the front a-priori Figure 1 DROPPED so the benchmark framework is now Figure 1 and the terminal descriptive
fit-matrix is Figure 8): the image is the nearest preceding image paragraph. This is robust to the figures
being grouped at the end OR placed inline between paragraphs.

Each image is embedded at THIS document's text-column width (page width minus L/R margins), so a figure
fills the column like the manuscript's own figures and never overflows the margins (A4 manuscript ≈ 6.27in;
standalone ≈ 6.00in). Height is recomputed from the fresh PNG aspect.

Run order: figure_*.py -> normalize_plate.py -> refresh_docx_figures.py -> update_docx_captions.py
           -> style_results_in_manuscript.py -> fix_table3.py -> clean_paper_temp.py
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image

WD = Path(__file__).resolve().parents[1]
PAPER = WD / "results" / "_paper"
BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
NUM_TO_KEY = {
    "2": "figure_ranking", "3": "figure_landscape",
    "4": "figure_cellcontext", "5": "figure_perturbation", "6": "figure_donor_decision",
    "7": "figure_immune_blindspot", "8": "figure_within_family_fit",
}


def figure_images(doc):
    """Map figure number ('1'..'8') -> the image paragraph that precedes its 'Figure N.' caption."""
    paras = doc.paragraphs
    out = {}
    for i, p in enumerate(paras):
        m = re.match(r"^Figure\s+([1-8])[.]", p.text.strip())
        if not m:
            continue
        for j in range(i - 1, -1, -1):                      # nearest preceding image paragraph
            if paras[j]._element.findall(".//" + BLIP):
                out.setdefault(m.group(1), paras[j])
                break
    return out


def refresh(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    sec = doc.sections[0]
    embed_w_in = (sec.page_width - sec.left_margin - sec.right_margin) / 914400.0
    imgs = figure_images(doc)
    missing = [n for n in NUM_TO_KEY if n not in imgs]
    if missing:
        sys.exit(f"FAIL {docx_path.name}: image(s) for Figure {missing} not found "
                 f"(each needs a 'Figure N.' caption right after its image)")
    for n, key in NUM_TO_KEY.items():
        para = imgs[n]
        png = PAPER / f"{key}.png"
        if not png.exists():
            sys.exit(f"FAIL: missing {png}")
        with Image.open(png) as im:
            w_px, h_px = im.size
        new_aspect = h_px / w_px
        for r in list(para.runs):                           # clear, then re-add at the column width
            r._element.getparent().remove(r._element)
        para.add_run().add_picture(str(png), width=Inches(embed_w_in))
        print(f"  {docx_path.name}: Figure {n} ({key:22s}) -> {embed_w_in:.2f} x "
              f"{embed_w_in*new_aspect:5.2f} in [{w_px}x{h_px}px]")
    doc.save(str(docx_path))
    print(f"  saved {docx_path.name}")


if __name__ == "__main__":
    paths = [Path(p) for p in sys.argv[1:]] or [
        PAPER / "results_section.docx",
        WD.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx",
    ]
    for p in paths:
        if not p.exists():
            sys.exit(f"FAIL: docx not found: {p}")
        print(f"refreshing figures in {p}")
        refresh(p)
    print("done.")
