#!/usr/bin/env python
"""Propagate the CellOT/chemCPA integration (Case 3 donor) into the deliverable manuscript .docx.

Body-prose: run-level find/replace (body paragraphs are single-run). Tables: add 2 Table-5 rows; rewrite
the Table-6 Soskic donor row; insert the § footnote. Figure legends are rebuilt separately by
update_docx_captions.py. Reports any unmatched pair so nothing is silently skipped.
"""
import shutil, sys
from pathlib import Path
from docx import Document
from docx.table import Table

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx"

PAIRS = [
 ("Eight conditioned models spanning six families were executed, each chosen for a maintained implementation that met a split's data requirements (Table 5); CINEMA-OT is carried as a perturbation-agnostic distributional comparator rather than a conditioned predictor.",
  "Ten conditioned models spanning seven families were executed, each chosen for a maintained implementation that met a split's data requirements (Table 5); CINEMA-OT is carried as a perturbation-agnostic distributional comparator rather than a conditioned predictor. The roster was extended to close two model-family coverage gaps, a neural optimal-transport predictor (CellOT) for stimulation and donor transfer and a native chemistry-aware compositional model (chemCPA) for unseen-compound prediction."),
 ("Third, model coverage is partial: eight conditioned models of the roughly thirty surveyed were executed",
  "Third, model coverage is partial: ten conditioned models of the roughly thirty surveyed were executed"),
 ("We therefore report OP3 as the clear cell-context result and Kang as competitive and local, and do not pool them.",
  "We therefore report OP3 as the clear cell-context result and Kang as competitive and local, and do not pool them. A neural optimal-transport predictor (CellOT), added as a task-matched stimulation-transfer model, was competitive on the same Kang split (mean conditioned-minus-baseline Pearson-Δ +0.012, cluster-bootstrap CI [−0.017, +0.045] over eight lineages, on par with scGen) while markedly improving distributional fidelity (lower E-distance on all eight lineages); it did not change the cell-context verdict (Supplementary Note S3)."),
 ("statistically equivalent to zero. Program-level recovery on the unseen-gene split was effectively degenerate",
  "statistically equivalent to zero. A native chemistry-aware compositional model evaluated on the same split, chemCPA with an RDKit-Morgan drug encoder, likewise did not exceed the no-chemistry baseline (0.11 versus 0.172) and predicted a near-constant response across the held-out compounds (Supplementary Note S4), so explicit molecular-structure conditioning did not change the unseen-compound verdict under the evaluated OP3 split. Program-level recovery on the unseen-gene split was effectively degenerate"),
 ("On the dedicated Soskic donor-resolved substrate, scGen, the conditioned model native to this activation contrast, did not exceed the pre-specified primary baseline under leave-one-donor-out: it averaged 0.144 against a per-donor primary baseline of 0.267, a donor-bootstrap gap of −0.123 [−0.137, −0.110] with 1.9% of donors positive (n = 106; Figure 6b, Table 6). The negative verdict was consistent across all three immune-aware metrics defined on Soskic (response-direction, distributional, and immune-program fidelity, each oriented so that positive favours conditioning; Supplementary). Kang is retained as a random-split optimism control rather than the main donor-axis evidence: random cell splits inflated Pearson-Δ by a small but suggestive margin relative to leak-proof leave-one-donor-out (Figure 6a).",
  "On the dedicated Soskic donor-resolved substrate the donor-axis verdict depended on the model family. The latent-shift model scGen, native to this activation contrast, did not exceed the pre-specified primary baseline under leave-one-donor-out: it averaged 0.144 against a per-donor primary baseline of 0.267, a donor-bootstrap gap of −0.123 [−0.137, −0.110] with 1.9% of donors positive (n = 106). A neural optimal-transport predictor, CellOT, on the same leak-proof split exceeded that baseline: it averaged 0.367 for a donor-bootstrap gap of +0.100 [+0.079, +0.121] with 87.7% of donors positive (p < 0.001), and the advantage held on all three immune-aware metrics (Figure 6b, Table 6; Supplementary Note S6). Donor-held-out activation transfer is therefore recoverable in this benchmark, but only by a model whose perturbation representation matches the distributional geometry of the response, so we no longer describe donor generalization as uniformly negative. Kang is retained as a random-split optimism control rather than the main donor-axis evidence: random cell splits inflated Pearson-Δ by a small but suggestive margin relative to leak-proof leave-one-donor-out (Figure 6a)."),
 ("Across the evaluated axes, perturbation conditioning exceeded the simple baselines only for seen-perturbation cell-context transfer, most clearly on OP3 and locally on Kang. It did not exceed them for unseen-gene, unseen-compound, or modality generalization in these splits, nor for the tested conditioned model (scGen) under Soskic leave-one-donor-out. Response-direction and immune-program fidelity also dissociated. These results support axis-stratified reporting rather than a pooled leaderboard.",
  "Across the evaluated axes, perturbation conditioning exceeded the simple baselines for seen-perturbation cell-context transfer (most clearly on OP3, locally on Kang) and, for a neural optimal-transport model, for donor-held-out activation transfer on Soskic. It did not exceed them for unseen-gene, unseen-compound, or modality generalization in these splits. The donor axis was model-family-dependent: the latent-shift scGen fell below the matched baseline under Soskic leave-one-donor-out, whereas the optimal-transport CellOT exceeded it, so donor-held-out performance depends on the perturbation-effect representation and must be evaluated under leak-proof donor splits. Response-direction and immune-program fidelity also dissociated. These results support axis-stratified reporting, and model-family-stratified reporting within the donor axis, rather than a pooled leaderboard."),
 ("but it did not exceed a non-conditioned training-mean or donor shift when the task was to extrapolate to an unseen perturbation, to stress modality, or to generalize to an unseen donor. Because the same model family changes verdict across axes,",
  "but it did not exceed a non-conditioned training-mean or donor shift when the task was to extrapolate to an unseen perturbation or to stress modality. On the donor axis the outcome depended on the model family: a latent-shift model did not exceed the baseline, whereas a neural optimal-transport model did. Because the same model family changes verdict across axes,"),
]

def all_runs(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            yield r
    for el in doc.element.body.iterchildren():
        if el.tag.split("}")[-1] == "tbl":
            for row in Table(el, doc).rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            yield r

def main():
    if not DOCX.exists():
        sys.exit(f"docx not found: {DOCX}")
    bak = DOCX.with_name(DOCX.stem + ".pre-models.docx")
    if not bak.exists():
        shutil.copy2(DOCX, bak)
    doc = Document(str(DOCX))

    # 1) body-prose run-level replacements
    hits = {o: 0 for o, _ in PAIRS}
    for r in all_runs(doc):
        for o, n in PAIRS:
            if o in r.text:
                r.text = r.text.replace(o, n); hits[o] += 1
    missed = [o[:50] for o, h in hits.items() if h == 0]
    print(f"body prose: {sum(hits.values())} applied" + (f"; NOT MATCHED: {missed}" if missed else "; all matched"))

    # 2) tables: locate Table 5 (Model|Family) and Table 6 (Axis|Dataset)
    t5 = t6 = None
    for t in doc.tables:
        hdr = [c.text.strip() for c in t.rows[0].cells]
        if hdr[:2] == ["Model", "Family"]:
            t5 = t
        if hdr and hdr[0] == "Axis":
            t6 = t

    if t5 is not None and not any("chemCPA" in r.cells[0].text for r in t5.rows):
        for cells in [["chemCPA", "latent (chemistry-aware)", "OP3 unseen compound", "native (RDKit-Morgan drug encoder)"],
                      ["CellOT", "optimal-transport", "Kang leave-one-lineage-out; Soskic leave-one-donor-out", "native (Kang) / adapted (Soskic)"]]:
            row = t5.add_row().cells
            for c, x in zip(row, cells):
                c.text = x
        print("Table 5: +2 rows (chemCPA, CellOT)")
    else:
        print("Table 5: rows already present or table not found")

    # 3) Table 6 Soskic donor row update
    if t6 is not None:
        for row in t6.rows:
            c0 = row.cells[0].text.strip()
            c1 = row.cells[1].text.strip()
            if c0 == "Donor" and c1.startswith("Soskic"):
                vals = ["Donor", "Soskic CD4 activation, leave-one-donor-out (0h->16h)", "donor (n = 106)",
                        "0.267 (best of training-mean/donor shift)", "CellOT 0.367 (scGen 0.144)", "—§",
                        "+0.100 [+0.079, +0.121] (CellOT); scGen −0.123 [−0.137, −0.110]",
                        "cluster bootstrap over donors; 87.7% positive (CellOT)",
                        "conditioning exceeds baseline (CellOT; latent-shift scGen below)"]
                for c, x in zip(row.cells, vals):
                    c.text = x
                print("Table 6: Soskic donor row updated to CellOT (Case 3)")
                break

    # 4) § footnote after the ‡ footnote
    if not any(p.text.strip().startswith("§ ") or p.text.strip().startswith("§\t") for p in doc.paragraphs):
        for p in doc.paragraphs:
            if p.text.strip().startswith("‡ Pooled descriptive values for the CRISPR row"):
                fn = ("§ The Soskic donor row reports the best applicable conditioned model (CellOT, optimal "
                      "transport); the latent-shift scGen is shown for comparison. CellOT's advantage is "
                      "corroborated by the other two immune-aware metrics (donor-bootstrap E-distance gap +4.95 "
                      "and AUCell-Δ gap +0.012, both p < 0.001, 100% and 92% of donors positive; Supplementary "
                      "Note S6). These are additional metric checks, not distributional comparators.")
                np_ = p.insert_paragraph_before(fn)
                np_.style = p.style
                if p.runs:
                    src = p.runs[0]
                    for rr in np_.runs:
                        rr.font.name = src.font.name; rr.font.size = src.font.size; rr.italic = src.italic
                # move it AFTER ‡ (insert_paragraph_before put it before ‡; swap)
                p._p.addnext(np_._p)
                print("Table 6: § footnote inserted")
                break

    doc.save(str(DOCX))
    print(f"saved {DOCX.name}; backup {bak.name}")

if __name__ == "__main__":
    main()
