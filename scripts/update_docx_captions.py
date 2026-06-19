#!/usr/bin/env python3
"""Replace the five figure-caption paragraphs in a results .docx IN PLACE from results_section.md.

Only the caption (figure-legend) paragraphs change — the body prose, headings, Table 3, and table
footnotes are left byte-for-byte intact. Each caption is matched in the docx by its "Fig. N |" prefix,
its runs are cleared, and new runs are rebuilt from the markdown caption (``**bold**`` -> bold run),
preserving the paragraph's style. Run AFTER refresh_docx_figures.py (order is independent, but keeping
both in the rebuild chain documents intent):

    figure_*.py -> normalize_plate.py -> refresh_docx_figures.py -> update_docx_captions.py
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from docx import Document

WD = Path(__file__).resolve().parents[1]
MD = WD / "results" / "_paper" / "results_section.md"
FIG_NS = ["2", "3", "4", "5", "6"]


def captions_from_md() -> dict[str, str]:
    """Pull the five caption lines out of the markdown. Accepts the manuscript form
    '**Figure N. ...**' and the legacy Nature form '**Fig. N | ...**'."""
    text = MD.read_text(encoding="utf-8")
    out: dict[str, str] = {}
    for line in text.splitlines():
        m = re.match(r"^\*\*(?:Figure\s+([2-6])\.|Fig\.\s+([2-6])\s+\|)", line.strip())
        if m:
            out[m.group(1) or m.group(2)] = line.strip()
    missing = [n for n in FIG_NS if n not in out]
    if missing:
        sys.exit(f"FAIL: captions for Fig {missing} not found in {MD}")
    return out


def rebuild_runs(paragraph, md_caption: str) -> None:
    """Clear the paragraph and re-add runs, making **...** segments bold."""
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    # split on ** ; segments alternate normal / bold / normal / ...
    for i, seg in enumerate(md_caption.split("**")):
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        run.bold = (i % 2 == 1)


def update(docx_path: Path, caps: dict[str, str]) -> None:
    doc = Document(str(docx_path))
    done = set()
    for para in doc.paragraphs:
        t = para.text.strip()
        for n in FIG_NS:
            if n in done:
                continue
            if t.startswith(f"Figure {n}.") or t.startswith(f"Fig. {n} |"):
                rebuild_runs(para, caps[n])
                done.add(n)
                print(f"  {docx_path.name}: updated Figure {n} caption ({len(caps[n])} chars)")
                break
    missing = [n for n in FIG_NS if n not in done]
    if missing:
        sys.exit(f"FAIL {docx_path.name}: caption paragraph(s) for Fig {missing} not found in docx")
    doc.save(str(docx_path))
    print(f"  saved {docx_path.name}")


if __name__ == "__main__":
    caps = captions_from_md()
    paths = [Path(p) for p in sys.argv[1:]] or [
        WD / "results" / "_paper" / "results_section.docx",
        WD.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx",
    ]
    for p in paths:
        if not p.exists():
            sys.exit(f"FAIL: docx not found: {p}")
        print(f"updating captions in {p}")
        update(p, caps)
    print("done.")
