#!/usr/bin/env python3
"""Propagate this round's Results prose / Table-3 text edits from the markdown into the .docx files.

The .docx body prose and Table-3 cells are static (only the figure captions auto-rebuild from the md via
update_docx_captions.py). So a small, explicit set of find/replace pairs carries the latest reviewer-fix
wording into both deliverables. Each `old` is matched inside a single run (body paragraphs and table cells
are single-run), so run-level replace is exact and preserves the run's font. Reports any pair that did not
match (so a stale/renamed string is caught, not silently skipped)."""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

WD = Path(__file__).resolve().parents[1]

# (old, new) — old must match the CURRENT docx text exactly (unicode minus "−" where used).
# This round (Soskic integration review): Results connectivity (P2), claim-scope narrowing (P3),
# multi-metric robustness sentence (P4), and "five datasets" phrasing (P1). Mirrors the same edits
# made in results_section.md so the md stays canonical and the docx body prose tracks it.
REPLACEMENTS = [
    # P2 connectivity: frame Results as a rigorous, gated instantiation (not a complete solution)
    ("This Results section instantiates that framework on the subset of the curated landscape",
     "This Results section instantiates that framework — not as a complete solution to the immune "
     "virtual-cell problem, but as a rigorous, reproducible benchmark on the subset of current public resources"),
    ("with broad method applicability (Figure 2).",
     "with computable simple baselines and broad method applicability (Figure 2)."),
    # P1 phrasing: "all five datasets" -> "all five benchmark components"
    ("are computed on all five datasets.",
     "are computed on all five benchmark components."),
    # P3 donor: name scGen as the tested conditioned model (not all conditioning)
    ("On the dedicated Soskic donor-resolved substrate, conditioning did not exceed the pre-specified primary baseline under leave-one-donor-out:",
     "On the dedicated Soskic donor-resolved substrate, the tested conditioned model — scGen, the conditioned model native to this activation contrast — did not exceed the pre-specified primary baseline under leave-one-donor-out:"),
    # P4 multi-metric robustness sentence (E-distance, AUCell-Δ MAE move the same way)
    ("with 1.9% donors positive (n = 106; Figure 6b, Table 6). Kang is therefore retained",
     "with 1.9% donors positive (n = 106; Figure 6b, Table 6). This negative verdict was consistent across "
     "all three immune-aware metrics applied to Soskic: with each oriented so that a positive value favours "
     "conditioning, the donor-bootstrap differences were −0.123 [−0.137, −0.110] for response direction "
     "(Pearson-Δ), −0.056 [−0.089, −0.027] for distributional fidelity (E-distance), and −0.012 "
     "[−0.017, −0.006] for immune-program recovery (AUCell-Δ MAE), all below zero (source data). Kang is "
     "therefore retained"),
    # P3 small-molecule: unseen-compound conclusion limited to the single OP3 resource
    ("Across both perturbation sub-axes, when the held-out intervention is novel the conditioned families",
     "Across both perturbation sub-axes — with the unseen-compound half resting on the single OP3 compound "
     "resource and therefore reported as a statement about that split — when the held-out intervention is "
     "novel the conditioned families"),
    # P3 modality: Frangieh is a stress test, not a comprehensive immune-cell modality benchmark
    ("the CRISPR perturbations target tumour cells, while the RNA and CITE-seq readouts include immune-relevant surface markers.",
     "the CRISPR perturbations target tumour cells, while the RNA and CITE-seq readouts include "
     "immune-relevant surface markers. Because the perturbed cells are tumour rather than immune cells and "
     "the surface readout is a 20-marker panel, we read Frangieh as a modality stress test, not as a "
     "comprehensive benchmark of direct immune-cell perturbation across modalities."),
    # P3 temporal: 40h/5d explicitly outside scope (limitation only)
    ("Soskic now satisfies this criterion for the 0h-versus-16h donor axis, whereas temporal saturation and several adjacent resources remain outside the present benchmark.",
     "Soskic now satisfies this criterion for the 0h-versus-16h donor axis, whereas its later activation "
     "timepoints (40 h, 5 d) and any temporal interpolation or extrapolation, together with several adjacent "
     "resources, remain outside the present benchmark and are not claimed."),
    # P3 synthesis donor: scGen-specific
    ("nor for donor generalization in the Soskic leave-one-donor-out analysis.",
     "nor for the tested conditioned model (scGen) under Soskic donor generalization (leave-one-donor-out)."),
]


def all_runs(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            yield r
    for el in doc.element.body.iterchildren():
        if el.tag.split("}")[-1] == "tbl":
            for row in Table(el, doc).rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            yield r


def apply(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    hits = {old: 0 for old, _ in REPLACEMENTS}
    for r in all_runs(doc):
        for old, new in REPLACEMENTS:
            if old in r.text:
                r.text = r.text.replace(old, new)
                hits[old] += 1
    doc.save(str(docx_path))
    missed = [old[:48] for old, n in hits.items() if n == 0]
    print(f"  {docx_path.name}: applied {sum(hits.values())} replacement(s)"
          + (f"; NOT MATCHED: {missed}" if missed else "; all pairs matched"))


if __name__ == "__main__":
    paths = [Path(p) for p in sys.argv[1:]] or [
        WD / "results" / "_paper" / "results_section.docx",
        WD.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx",
    ]
    for p in paths:
        if not p.exists():
            sys.exit(f"FAIL: docx not found: {p}")
        print(f"syncing edits into {p}")
        apply(p)
    print("done.")
