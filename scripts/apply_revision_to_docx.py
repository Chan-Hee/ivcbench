#!/usr/bin/env python
"""Propagate the reviewer-revision edits into the deliverable manuscript .docx.

Body/caption/footnote = run-level find/replace; Table 3 and Table 6 cells = direct cell edits.
Figure legends are rebuilt separately by update_docx_captions.py from the md. Reports unmatched pairs.
"""
import shutil, sys
from pathlib import Path
from docx import Document
from docx.table import Table

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx"

PAIRS = [
 # methods-funnel count harmonization
 ("Ten conditioned models spanning seven families were executed, each chosen for a maintained implementation that met a split's data requirements (Table 5); CINEMA-OT is carried as a perturbation-agnostic distributional comparator rather than a conditioned predictor. The roster was extended to close two model-family coverage gaps, a neural optimal-transport predictor (CellOT) for stimulation and donor transfer and a native chemistry-aware compositional model (chemCPA) for unseen-compound prediction.",
  "Ten conditioned entrants were executed across the surveyed model classes, each chosen for a maintained implementation that met a split's data requirements (Table 5); CINEMA-OT is carried as a perturbation-agnostic distributional comparator rather than a conditioned predictor. CellOT and chemCPA were added to close optimal-transport and chemistry-aware coverage gaps, for stimulation and donor transfer and for unseen-compound prediction respectively."),
 # Limitations: statistical power -> effective replication
 ("Several limitations bound these conclusions. First, statistical power is modest: the resampling unit is the biological replicate (6 immune lineages, 5 datasets, 106 Soskic donors, or 8 Kang donors for the random-split control), so the axis-level endpoints are consistent-direction effects rather than high-powered tests.",
  "Several limitations bound these conclusions. First, effective replication varies by axis: OP3 cell-context and primary-T CRISPR are limited by lineage- or dataset-level replication (6 immune lineages, 5 datasets), whereas the Soskic donor axis has broader donor-level replication (106 donors) but tests a single 0 h to 16 h activation contrast; the axis-level endpoints are therefore consistent-direction effects rather than high-powered tests."),
 # Limitations: single positive axis -> scope-aware
 ("Fourth, the single positive axis (seen-compound cell-context transfer on OP3) is also the comparatively easier task, interpolating an already-seen perturbation across lineages, and should not be read as general evidence that conditioning transfers.",
  "Fourth, the positive settings have different scopes: OP3 cell-context transfer interpolates seen compounds across immune lineages, whereas Soskic donor transfer is a single 0 h to 16 h activation contrast in which CellOT, but not scGen, exceeded the matched baseline; neither should be read as general evidence that conditioning transfers."),
 # donor conclusion scope
 ("Donor-held-out activation transfer is therefore recoverable in this benchmark, but only by a model whose perturbation representation matches the distributional geometry of the response, so we no longer describe donor generalization as uniformly negative.",
  "Donor-held-out activation transfer was therefore recoverable for CellOT under the evaluated Soskic 0 h to 16 h split, by a model whose perturbation representation matches the distributional geometry of the response, so we no longer describe donor generalization as uniformly negative."),
 # Table 6 caption
 ("Statistical-decision table: native conditioning versus the pre-specified simple baseline",
  "Statistical-decision table: best applicable conditioning versus the pre-specified simple baseline"),
 # Table 6 footnote (section sign)
 ("The Soskic donor row reports the best applicable conditioned model (CellOT, optimal transport); the latent-shift scGen is shown for comparison.",
  "CellOT is adapted on Soskic but is the best applicable conditioned model for the donor axis; scGen is shown as the latent state-transfer comparator (native/adapted labels are defined in Table 5)."),
]

def all_runs(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            yield r
    for el in doc.element.body.iterchildren():
        if el.tag.split("}")[-1] == "tbl":
            for row in Table(el, doc).rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            yield r

def main():
    bak = DOCX.with_name(DOCX.stem + ".pre-revision.docx")
    if not bak.exists():
        shutil.copy2(DOCX, bak)
    doc = Document(str(DOCX))

    hits = {o: 0 for o, _ in PAIRS}
    for r in all_runs(doc):
        for o, n in PAIRS:
            if o in r.text:
                r.text = r.text.replace(o, n); hits[o] += 1
    missed = [o[:45] for o, h in hits.items() if h == 0]
    print(f"run-level: {sum(hits.values())} applied" + (f"; NOT MATCHED: {missed}" if missed else "; all matched"))

    # locate tables
    t3 = t6 = None
    for t in doc.tables:
        h = [c.text.strip() for c in t.rows[0].cells]
        if h and h[0].startswith("Task cluster"):
            t3 = t
        if h and h[0] == "Axis":
            t6 = t

    # Table 3: Kang + Soskic axis cells
    if t3 is not None:
        for row in t3.rows:
            c0 = row.cells[0].text.strip()
            if c0.startswith("Cytokine-response"):
                row.cells[2].text = "Cell-context transfer (leave-one-lineage-out); random-split donor optimism control"
            elif c0.startswith("Donor-aware activation"):
                row.cells[2].text = "Primary donor axis (leave-one-donor-out)"
        print("Table 3: Kang + Soskic axis cells updated")

    # Table 6: header cell + Soskic row 0h->16h
    if t6 is not None:
        for c in t6.rows[0].cells:
            if c.text.strip() == "Best native conditioned":
                c.text = "Best applicable conditioned model"
        for row in t6.rows:
            if row.cells[0].text.strip() == "Donor" and "Soskic" in row.cells[1].text:
                row.cells[1].text = row.cells[1].text.replace("0h->16h", "0 h to 16 h")
        print("Table 6: header + Soskic 0 h to 16 h updated")

    doc.save(str(DOCX))
    print(f"saved {DOCX.name}; backup {bak.name}")

if __name__ == "__main__":
    main()
