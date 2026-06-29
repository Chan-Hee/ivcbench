#!/usr/bin/env python3
"""Propagate the finalized v2 Results + Discussion prose into the deliverable manuscript .docx.

This is the structural propagation the small in-place edit scripts (apply_loop_revisions_to_docx.py,
update_docx_captions.py, style_results_in_manuscript.py) cannot do, because v2 is a full restructure:
  * the OLD benchmark Results prose (under "An Immune-Aware Benchmark for Perturbation Prediction") is
    REPLACED by results_section_v2.md prose;
  * the Discussion (discussion_v2.md) is INSERTED after the Results region;
  * the figure plate goes from the OLD 6-figure scheme (front a-priori Figure 1 + benchmark Figs 2-6) to
    the v2 8-figure scheme (front Fig 1 DROPPED; benchmark Figs 1-8), so 8 image+caption paragraphs are
    laid out in order 1..8 for refresh_docx_figures.py to populate;
  * the front a-priori Figure 1 (image + caption) is DROPPED and its one referencing sentence reframed.

Tables 3/4/5/6 (captions + footnotes) are PRESERVED verbatim and re-threaded into the v2 layout at the
points the v2 prose references them. Image paragraphs are left empty here (placeholders) — the figure
rebuild chain (normalize_plate -> refresh_docx_figures -> update_docx_captions[v2] -> style) fills them.

Run order:  apply_v2_to_docx.py  ->  normalize_plate.py  ->  refresh_docx_figures.py
            ->  (captions are written by THIS script directly, so update_docx_captions is optional)
            ->  style_results_in_manuscript_v2 (font normalize)  ->  clean_paper_temp.py

Idempotence: detects whether v2 was already applied (a Discussion heading present) and refuses to double-apply.
Every number in the inserted prose comes verbatim from results_section_v2.md / discussion_v2.md; this script
transcribes, it does not compute.
"""
from __future__ import annotations
import copy
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "results" / "_paper"
RES_MD = PAPER / "results_section_v2.md"
DISC_MD = PAPER / "discussion_v2.md"
DOCX = ROOT.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx"

FONT = "Times New Roman"
BODY_PT, SUBHEAD_PT, SECTION_PT, CAPTION_PT = 10.0, 12.0, 16.0, 9.0
BLIP = ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"

SEC_HEAD = "An Immune-Aware Benchmark for Perturbation Prediction"
POST_MARKER = "Toward Immune Virtual Cells"

# ---- v2 figure caption texts (renumber map figure_renumber_final.md). Figs 1,3,4,5,6 reuse the approved
# OLD captions (renumbered); Fig 2 (ranking) uses the v2 cross-ref one-liner; Figs 7,8 are the v2 inline
# captions. Each is markdown with **bold** lead; pulled live from the deposited md where present.
def _old_caption(md_text: str, old_num: str) -> str:
    for line in md_text.splitlines():
        if line.strip().startswith(f"**Figure {old_num}."):
            return line.strip()
    sys.exit(f"FAIL: old caption Figure {old_num} not found")

def _v2_caption(md_text: str, num: str) -> str:
    for line in md_text.splitlines():
        if line.strip().startswith(f"**Figure {num}."):
            return line.strip()
    sys.exit(f"FAIL: v2 caption Figure {num} not found")


def build_captions() -> dict[str, str]:
    old = (PAPER / "results_section.md").read_text(encoding="utf-8")
    v2 = RES_MD.read_text(encoding="utf-8")
    caps = {}
    # v2 Fig 1 = framework (old Fig 2); 3=landscape(old 3); 4=cellcontext(old 4); 5=perturbation(old 5);
    # 6=donor(old 6). Renumber the leading "**Figure N." token.
    caps["1"] = re.sub(r"^\*\*Figure 2\.", "**Figure 1.", _old_caption(old, "2"))
    caps["3"] = _old_caption(old, "3")            # already "Figure 3"
    caps["4"] = _old_caption(old, "4")
    caps["5"] = _old_caption(old, "5")
    caps["6"] = _old_caption(old, "6")
    # Fig 2 = ranking — no full caption exists; synthesize from the v2 cross-ref descriptor.
    caps["2"] = ("**Figure 2. Method × split ranking heatmap.** Ranking heatmap of response-direction "
                 "Pearson-Δ for every applicable (model, split) pair across the benchmark, rows grouped "
                 "by method family and columns by evaluation split, on a zero-centred diverging scale "
                 "(blues positive, rust negative). This is the scPerturBench-style whole-panel ranking view "
                 "of the same values adjudicated against the universal floor in Figure 3 and Table 6; "
                 "schematic ordering, no error bars.")
    caps["7"] = _v2_caption(v2, "7")
    caps["8"] = _v2_caption(v2, "8")
    return caps


# ------------------------- markdown inline -> docx runs -------------------------
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")

def _clean(seg: str) -> str:
    # strip backticks; drop markdown link/footnote scaffolding but keep readable text
    return seg

def add_inline(p: Paragraph, text: str, base_pt: float, lead_bold: bool = False):
    """Add runs to paragraph p, honouring **bold**, *italic*, `code`(plain). base_pt sets size."""
    # Normalise footnote anchors like [^c5n] -> drop; keep the footnote body lines handled separately.
    text = re.sub(r"\[\^[^\]]+\]", "", text)
    parts = INLINE.split(text)
    first_nonblank = True
    for seg in parts:
        if seg == "":
            continue
        bold = italic = False
        s = seg
        if seg.startswith("**") and seg.endswith("**"):
            s = seg[2:-2]; bold = True
        elif seg.startswith("*") and seg.endswith("*") and len(seg) > 1:
            s = seg[1:-1]; italic = True
        elif seg.startswith("`") and seg.endswith("`"):
            s = seg[1:-1]
        if s == "":
            continue
        r = p.add_run(s)
        r.font.name = FONT
        r.font.size = Pt(base_pt)
        r.bold = bold or (lead_bold and first_nonblank)
        r.italic = italic
        first_nonblank = False


def new_para_like(anchor: Paragraph) -> Paragraph:
    """Clone an existing body paragraph's <w:p> (to inherit pPr/style), clear its runs, return Paragraph."""
    new_p = copy.deepcopy(anchor._p)
    # remove all runs
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    return Paragraph(new_p, anchor._parent)


# ------------------------- markdown block parser -------------------------
def parse_blocks(md_text: str):
    """Yield (kind, text) blocks: kind in {h1,h2,body,figcap,footnote}. Skips hr/em-only meta lines."""
    lines = md_text.splitlines()
    out = []
    for raw in lines:
        line = raw.rstrip()
        s = line.strip()
        if not s:
            continue
        if s == "---":
            continue
        if s.startswith("# "):
            out.append(("h1", s[2:].strip())); continue
        if s.startswith("## "):
            out.append(("h2", s[3:].strip())); continue
        if s.startswith("### "):
            out.append(("h3", s[4:].strip())); continue
        if s.startswith("**Figure ") :
            out.append(("figcap", s)); continue
        if s.startswith("[^"):              # footnote definition
            out.append(("footnote", re.sub(r"^\[\^[^\]]+\]:\s*", "", s))); continue
        if s.startswith("> "):              # blockquote -> body (keep content)
            out.append(("body", s[2:].strip())); continue
        if s.startswith("- "):              # list item -> body with bullet
            out.append(("li", s[2:].strip())); continue
        # the leading *italic meta* paragraph (provenance note) -> keep as italic body
        out.append(("body", s))
    return out


def main():
    doc = Document(str(DOCX))
    blocks = list(doc.element.body.iterchildren())

    # idempotence guard
    for el in blocks:
        if el.tag.split("}")[-1] == "p":
            if Paragraph(el, doc).text.strip().startswith("Discussion: what the immune blind-spot map"):
                sys.exit("REFUSE: v2 already applied (Discussion heading present). Restore the .pre-v2 backup first.")

    # locate region indices (body-block indices)
    def find_block(pred):
        for i, el in enumerate(blocks):
            if el.tag.split("}")[-1] == "p" and pred(Paragraph(el, doc).text.strip()):
                return i
        return None

    i_sec = find_block(lambda t: t == SEC_HEAD)
    i_post = find_block(lambda t: t == POST_MARKER)
    if i_sec is None or i_post is None:
        sys.exit(f"FAIL: section/post markers not found (sec={i_sec}, post={i_post})")

    # collect the TABLE blocks (caption P + tbl) and Table-6 footnote Ps to preserve, by scanning region
    def text_of(i):
        el = blocks[i]
        return Paragraph(el, doc).text.strip() if el.tag.split("}")[-1] == "p" else "<TBL>"

    region = list(range(i_sec, i_post))
    # identify preserve sets
    tbl_caption_starts = {"Table 3.": "T3", "Table 4.": "T4", "Table 5.": "T5", "Table 6.": "T6"}
    preserve = {}   # tag -> list of cloned elements (caption P, tbl, [footnotes for T6])
    i = i_sec
    while i < i_post:
        t = text_of(i)
        hit = next((tag for pre, tag in tbl_caption_starts.items() if t.startswith(pre)), None)
        if hit:
            grp = [copy.deepcopy(blocks[i])]            # caption paragraph
            j = i + 1
            # next block should be the table
            if blocks[j].tag.split("}")[-1] == "tbl":
                grp.append(copy.deepcopy(blocks[j])); j += 1
            # Table 6: also grab its baseline-legend + footnote paragraphs (†/‡/§/¶ and 'Simple baselines')
            if hit == "T6":
                while j < i_post:
                    tj = text_of(j)
                    if tj.startswith(("Simple baselines", "†", "‡", "§", "¶")):
                        grp.append(copy.deepcopy(blocks[j])); j += 1
                    else:
                        break
            preserve[hit] = grp
            i = j
        else:
            i += 1
    for tag in ("T3", "T4", "T5", "T6"):
        if tag not in preserve:
            sys.exit(f"FAIL: could not capture {tag} for preservation")

    # a reference body paragraph to clone style from (block i_sec+2 is a body para)
    body_ref = Paragraph(copy.deepcopy(blocks[i_sec + 2]), doc)

    caps = build_captions()

    # -------- build the new Results+Discussion element stream --------
    def mk_para(kind, text, table_tag=None):
        p = new_para_like(body_ref)
        if kind == "h1":      # section heading already present (we keep the existing one) -> skip
            add_inline(p, text, SECTION_PT, lead_bold=False);
            for r in p.runs: r.bold = True
        elif kind == "h2":
            add_inline(p, text, SUBHEAD_PT);
            for r in p.runs: r.bold = True
        elif kind == "h3":
            add_inline(p, text, SUBHEAD_PT)
            for r in p.runs: r.bold = True
        elif kind == "figcap":
            add_inline(p, text, CAPTION_PT, lead_bold=True)
        elif kind == "footnote":
            add_inline(p, text, CAPTION_PT)
            for r in p.runs: r.italic = True
        elif kind == "li":
            add_inline(p, "• " + text, BODY_PT)
        else:
            add_inline(p, text, BODY_PT)
        return p._p

    # embed-width = text column of THIS document (page width minus L/R margins), matching the host figures
    sec0 = doc.sections[0]
    EMBED_W = (sec0.page_width - sec0.left_margin - sec0.right_margin)

    def emit_fig(num):
        """Emit an image paragraph (embedded PNG) + its caption paragraph, in numeric order."""
        from docx.shared import Emu
        key = {"2": "figure_ranking", "3": "figure_landscape",
               "4": "figure_cellcontext", "5": "figure_perturbation", "6": "figure_donor_decision",
               "7": "figure_immune_blindspot", "8": "figure_within_family_fit"}[num]
        png = PAPER / f"{key}.png"
        if not png.exists():
            sys.exit(f"FAIL: figure PNG missing for Figure {num}: {png}")
        ip = new_para_like(body_ref)
        ip.add_run().add_picture(str(png), width=Emu(int(EMBED_W)))
        new_stream.append(ip._p)
        new_stream.append(mk_para("figcap", caps[num]))

    new_stream = []   # list of oxml elements to insert after the section heading

    # --- Results v2 prose ---
    res_blocks = parse_blocks(RES_MD.read_text(encoding="utf-8"))
    # drop the leading h1 (== section heading, already in doc)
    res_blocks = [b for b in res_blocks if not (b[0] == "h1")]
    # drop the inline Fig 7 / Fig 8 caption blocks from the prose flow: we emit ALL eight captions in
    # strict numeric order (1..8) alongside their embedded images, so document order == figure number.
    res_blocks = [b for b in res_blocks if not (b[0] == "figcap")]

    # Figure placement, threaded so caption numbers ASCEND monotonically (1..8) in document order:
    #   subhead -> §1 prose -> Fig 1 + Tables 3/4/5 -> §2 prose -> Fig 2, Fig 3 -> §3 prose -> Fig 4,5,6
    #   -> §4 prose -> §5 prose -> Fig 7, Fig 8 -> cross-refs -> Table 6
    FIG_AT_SECTION_END = {            # emit these figures AFTER the section whose number is the key finishes
        "1.": ["1"],                  # design charter -> Fig 1 framework (+ Tables 3/4/5)
        "2.": ["2", "3"],             # coverage/floor -> ranking + landscape
        "3.": ["4", "5", "6"],        # dissociation -> cellcontext + perturbation + donor
        "4.": [],                     # within-family
        "5.": ["7", "8"],             # blind-spot -> immune_blindspot (7) + within_family_fit (8)
    }

    # First subhead
    new_stream.append(mk_para("h3", "Benchmark design and scope"))

    cur_section = None
    pending_tables_after_fig1 = True
    for kind, text in res_blocks:
        if kind == "h2":
            if cur_section is not None:
                for fn in FIG_AT_SECTION_END.get(cur_section, []):
                    emit_fig(fn)
                    if fn == "1" and pending_tables_after_fig1:
                        for tg in ("T3", "T4", "T5"):
                            for el in preserve[tg]:
                                new_stream.append(el)
                        pending_tables_after_fig1 = False
            m = re.match(r"^(\d+)\.", text)
            cur_section = (m.group(1) + ".") if m else None
            new_stream.append(mk_para("h2", text))
        else:
            new_stream.append(mk_para(kind, text))

    # flush figures owed by the final section (§5 -> Fig 7, Fig 8)
    if cur_section is not None:
        for fn in FIG_AT_SECTION_END.get(cur_section, []):
            emit_fig(fn)

    # Table 6 goes at the end of the cross-references block (v2 names Table 6 as the decision table)
    for el in preserve["T6"]:
        new_stream.append(el)

    # --- Discussion v2 ---
    disc_blocks = parse_blocks(DISC_MD.read_text(encoding="utf-8"))
    # The discussion md h1 is the Discussion title; turn it into a section-level heading. Also it contains a
    # second h1 "Abstract-alignment..." which is editorial scaffolding (sentences the abstract must carry) ->
    # keep as a subheading block (informative, deposited), not deleted.
    for kind, text in disc_blocks:
        if kind == "h1":
            if text.startswith("Discussion"):
                new_stream.append(mk_para("h1", "Discussion: what the immune blind-spot map shows, and what it means for the fit-matrix"))
            else:
                new_stream.append(mk_para("h2", text))
        else:
            new_stream.append(mk_para(kind, text))

    # -------- splice: delete old region body (keep heading at i_sec), insert new_stream after heading --------
    body = doc.element.body
    sec_el = blocks[i_sec]            # the section heading paragraph (keep)
    # remove everything from i_sec+1 .. i_post-1 (the old Results body + tables + figs)
    for el in blocks[i_sec + 1:i_post]:
        body.remove(el)
    # insert new_stream right after the section heading, in order
    anchor = sec_el
    for el in new_stream:
        anchor.addnext(el)
        anchor = el

    # -------- drop the front a-priori Figure 1 (image + caption) and reframe its referencing sentence --------
    # re-fetch blocks (we mutated the tree)
    blocks2 = list(doc.element.body.iterchildren())
    fig1_img = fig1_cap = ref_sentence = None
    for el in blocks2:
        if el.tag.split("}")[-1] != "p":
            continue
        p = Paragraph(el, doc)
        t = p.text.strip()
        if t.startswith("Figure 1. Immune prediction tasks and corresponding families"):
            fig1_cap = el
        if el.findall(BLIP) and not t:
            # candidate image; only the FRONT a-priori one sits right before the Fig1 caption
            pass
        if "Figure 1 visualizes this correspondence as a 5×6 matrix" in t:
            ref_sentence = p
    # the front image is the image paragraph immediately preceding the Fig1 caption
    if fig1_cap is not None:
        prev = fig1_cap.getprevious()
        while prev is not None and prev.tag.split("}")[-1] == "p" and not prev.findall(BLIP) and Paragraph(prev, doc).text.strip() == "":
            prev = prev.getprevious()
        if prev is not None and prev.findall(BLIP):
            fig1_img = prev
    dropped = []
    if fig1_cap is not None:
        doc.element.body.remove(fig1_cap); dropped.append("front-Fig1-caption")
    if fig1_img is not None:
        doc.element.body.remove(fig1_img); dropped.append("front-Fig1-image")
    if ref_sentence is not None:
        # reframe: remove the "Figure 1 visualizes..." clause, point to the closing fit-matrix (Figure 8)
        new_text = ("With the six families in view, we ask which of them should most naturally align with "
                    "each of the five immune prediction tasks defined above. Below we set out the expected "
                    "suitability of each family per task, reasoning from the kind of perturbation effect each "
                    "family represents; this mapping is an a-priori map — a set of predictions, not yet a "
                    "result — and it is intentionally many-to-many, as several families could contribute "
                    "to multiple tasks in distinct capacities. We test it empirically in the benchmark "
                    "(Results), where each expected fit is confronted with real immune data under leak-safe "
                    "splits and a single universal simple-baseline floor; the evidence-based outcome is "
                    "summarised in the closing fit-recommendation matrix (Figure 8).")
        for r in list(ref_sentence.runs):
            r._element.getparent().remove(r._element)
        rr = ref_sentence.add_run(new_text); rr.font.name = FONT; rr.font.size = Pt(BODY_PT)
        dropped.append("front-ref-reframed")

    doc.save(str(DOCX))
    print("APPLIED v2.")
    print("  preserved tables:", ", ".join(sorted(preserve)))
    print("  front Fig 1 drop:", ", ".join(dropped) or "NONE (already gone?)")
    print("  inserted blocks:", len(new_stream))


if __name__ == "__main__":
    main()
