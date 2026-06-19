#!/usr/bin/env python
"""Propagate the review-loop's verified revisions from canonical results_section.md into the main .docx.
Body prose paragraphs are single-run -> run-level find/replace (round0 text -> canonical text), auto-derived
from the md diff (no hand transcription). Table-6 cell edits use EXACT round0->canonical substrings (bold
gap number preserved by per-run replace). The two new footnotes are cloned from the section footnote so they
inherit its italic 9pt style. Figure captions are left to update_docx_captions. Reports any unmatched item.
"""
import difflib, shutil, sys, copy
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
OLD_MD = ROOT / "outputs/review_loop/round0_main.snapshot.md"
NEW_MD = ROOT / "results/_paper/results_section.md"
DOCX = ROOT.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx"

old_lines = [l.rstrip("\n") for l in OLD_MD.open()]
new_lines = [l.rstrip("\n") for l in NEW_MD.open()]


def kind(l):
    if l.startswith("**Figure"):
        return "caption"
    if l.startswith("|"):
        return "table"
    if l.startswith("#") or not l.strip():
        return "skip"
    if l.startswith("*"):
        return "footnote"
    return "prose"


# auto prose/footnote pairs from the md diff (single-line replaces only)
prose_pairs = []
sm = difflib.SequenceMatcher(a=old_lines, b=new_lines, autojunk=False)
for tag, i1, i2, j1, j2 in sm.get_opcodes():
    if tag != "replace":
        continue
    olds = [l for l in old_lines[i1:i2] if kind(l) not in ("skip",)]
    news = [l for l in new_lines[j1:j2] if kind(l) not in ("skip",)]
    for o, n in zip(olds, news):
        if kind(o) in ("caption", "table") or kind(n) in ("caption", "table"):
            continue
        prose_pairs.append((o.lstrip("*").strip() if kind(o) == "footnote" else o,
                            n.lstrip("*").strip() if kind(n) == "footnote" else n))

# EXACT Table-6 cell edits (round0 substring -> canonical substring); bold gap number preserved per-run
CELL = [
    ("scGen 0.917 (seed 0)", "scGen 0.884 (3-seed mean)"),
    ("+0.098", "+0.065"),
    (" (seed range +0.02..+0.10; a seed range, not a cluster CI)",
     " mean (3-seed range +0.016 to +0.092; a seed range, not a cluster CI; CI crosses zero, under-powered (n = 3 seeds))"),
    ("3/3 seeds sign-robust", "3/3 seeds sign-robust, not a stable +0.10"),
    ("exceeds locally", "exceeds locally (under-powered)"),
    ("baseline not exceeded (severe)", "directional result (no cluster CI; seed range only): baseline not exceeded (severe)"),
]
# the two new footnotes (verbatim from canonical, asterisks stripped)
FOOTNOTES = [ln.lstrip("*").rstrip("*").strip() for ln in new_lines if ln.startswith("*¶")]


def all_paras(doc):
    for p in doc.paragraphs:
        yield p
    for el in doc.element.body.iterchildren():
        if el.tag.split("}")[-1] == "tbl":
            for row in Table(el, doc).rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        yield p


def run_replace(doc, find, repl, once=True):
    hits = 0
    for p in all_paras(doc):
        for run in p.runs:
            if find in run.text:
                run.text = run.text.replace(find, repl)
                hits += 1
                if once:
                    return hits
    if hits == 0:  # cross-run fallback (plain paragraphs)
        for p in all_paras(doc):
            joined = "".join(r.text for r in p.runs)
            if find in joined and p.runs:
                p.runs[0].text = joined.replace(find, repl)
                for r in p.runs[1:]:
                    r.text = ""
                return 1
    return hits


def main():
    doc = Document(str(DOCX))
    miss = []

    for o, n in prose_pairs:
        h = run_replace(doc, o, n)
        if h == 0:
            miss.append("PROSE:" + o[:50])
    print(f"prose: {len(prose_pairs)} pairs, {len(miss)} unmatched")

    for fo, fn in CELL:
        if run_replace(doc, fo, fn) == 0:
            miss.append("CELL:" + fo[:40])

    # Frangieh RNA verdict (row-targeted; 'baseline not exceeded' is not unique)
    done_rna = False
    for el in doc.element.body.iterchildren():
        if el.tag.split("}")[-1] == "tbl":
            for row in Table(el, doc).rows:
                texts = [c.text for c in row.cells]
                if any("Frangieh, unseen KO (RNA)" in t for t in texts):
                    vc = row.cells[-1]
                    if vc.text.strip() == "baseline not exceeded":
                        for p in vc.paragraphs:
                            if p.runs:
                                p.runs[0].text = "directional result (no cluster CI; seed range only): baseline not exceeded"
                                for r in p.runs[1:]:
                                    r.text = ""
                        done_rna = True
    if not done_rna:
        miss.append("CELL:Frangieh-RNA-verdict")

    # insert the two new footnotes after the section (§) footnote, cloning its style
    sec = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("§ CellOT is adapted on Soskic") or "CellOT is adapted on Soskic but is the best applicable" in p.text:
            sec = p
    if sec is None:
        miss.append("INSERT:section-footnote-anchor")
    else:
        anchor = sec
        for fn in FOOTNOTES:
            new_p = copy.deepcopy(anchor._p)
            anchor._p.addnext(new_p)
            np = Paragraph(new_p, anchor._parent)
            for r in list(np.runs):
                r.text = ""
            (np.runs[0] if np.runs else np.add_run("")).text = fn
            anchor = np
        print(f"inserted {len(FOOTNOTES)} footnote(s) after section footnote")

    if miss:
        print("!! UNMATCHED: " + " | ".join(miss), file=sys.stderr)
    doc.save(str(DOCX))
    print("saved" + ("" if not miss else f" WITH {len(miss)} MISS"))
    sys.exit(2 if miss else 0)


if __name__ == "__main__":
    main()
