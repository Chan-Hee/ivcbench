#!/usr/bin/env python
"""Propagate the verified CellOT 3-seed Soskic numbers + the one metric sentence into the deliverable
.docx. Per-run replace (preserves the Table 6 bold gap cell); cross-run collapse fallback for any pair
that spans runs (plain paragraphs only). Figure 6 legend is re-synced separately by update_docx_captions.
"""
import os, shutil, sys
from docx import Document
from docx.table import Table

DOCX = os.environ.get(
    "IVCBENCH_MANUSCRIPT_DOCX",
    "Toward Immune Virtual Cells (benchmark Results inserted).docx",
)  # set $IVCBENCH_MANUSCRIPT_DOCX to the manuscript path (not shipped in this release)
BAK = DOCX.replace(".docx", ".pre-3seed.docx")

PAIRS = [
    ("+0.100", "+0.102"),
    ("[+0.079, +0.121]", "[+0.083, +0.122]"),
    ("averaged 0.367", "averaged 0.369"),
    ("CellOT 0.367", "CellOT 0.369"),
    ("E-distance gap +4.95", "E-distance gap +4.94"),
    ("100% and 92% of donors positive", "100% and 96.2% of donors positive"),
    ("87.7% positive (CellOT)", "87.7% positive (CellOT, seeds 0–2)"),
    ("(p < 0.001), and the advantage held on all three",
     "(p < 0.001; across three training seeds, mean ±0.004), and the advantage held on all three"),
    ("every component except Frangieh (Table 4).",
     "every component except Frangieh (Table 4). Pearson-Δ scores response direction and is "
     "amplitude-invariant by design, restricted to response genes because all-gene correlation is "
     "inflated by the unperturbed majority, while the paired E-distance supplies the magnitude and "
     "calibration axis."),
]


def paras(doc):
    for p in doc.paragraphs:
        yield p
    for el in doc.element.body.iterchildren():
        if el.tag.split("}")[-1] == "tbl":
            for row in Table(el, doc).rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        yield p


def main():
    if not os.path.exists(BAK):
        shutil.copy2(DOCX, BAK)
    d = Document(DOCX)
    hits = {f: 0 for f, _ in PAIRS}
    # per-run (preserves formatting)
    for p in paras(d):
        for run in p.runs:
            for f, r in PAIRS:
                if f in run.text:
                    run.text = run.text.replace(f, r)
                    hits[f] += 1
    # cross-run fallback for spanning pairs (plain paragraphs)
    for p in paras(d):
        joined = "".join(run.text for run in p.runs)
        for f, r in PAIRS:
            if hits[f] == 0 and f in joined and p.runs:
                p.runs[0].text = joined.replace(f, r)
                for run in p.runs[1:]:
                    run.text = ""
                hits[f] += 1
                joined = "".join(run.text for run in p.runs)
    for f, h in hits.items():
        print(f"  {h}x  {f[:46]!r}")
    miss = [f[:46] for f, h in hits.items() if h == 0]
    if miss:
        print(f"!! NOT MATCHED: {miss}", file=sys.stderr)
        sys.exit(1)
    d.save(DOCX)
    print(f"saved; backup {os.path.basename(BAK)}")


if __name__ == "__main__":
    main()
