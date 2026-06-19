#!/usr/bin/env python
"""Propagate the reviewer-disclosure edits into the deliverable manuscript .docx.

Four body/cell/footnote find-replace pairs (CPA/STATE cap note, 40-epoch note, Frangieh seed-range
qualifier, CellOT per-fold adapted/baseline note). Cross-run aware: if the find string spans runs in a
paragraph or table cell, the paragraph's runs are collapsed into the first run (formatting is uniform in
each target paragraph). Backs up to .pre-disclosures.docx and reports hit counts; fails loudly on a miss.
"""
import shutil, sys
from pathlib import Path
from docx import Document
from docx.table import Table

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx"

PAIRS = [
 # CPA/STATE cap disclosure (body methods paragraph)
 ("are not reported as native capability. On unseen-gene",
  "are not reported as native capability. To bound wall-clock on the largest panels, CPA and STATE "
  "train under a stratified cell cap (60,000 and 50,000 cells respectively), whereas scGen, CellOT, "
  "FP-ridge, and the simple baselines use the full training split; the cap is verdict-invariant, being "
  "inert on the Kang cell-context split and, on the CRISPR gene holdout, leaving the capped CPA above "
  "the uncapped scGen while both remain below the training-mean baseline. On unseen-gene"),
 # 40-epoch documentation (modality paragraph)
 ("and a wrong-sign exception on PD-L1 (CD274).",
  "and a wrong-sign exception on PD-L1 (CD274). scGen was trained for 40 epochs on this axis versus 60 "
  "on the others, for wall-clock; this is not load-bearing, because the epoch-free linear "
  "knockout-embedding shift reproduces the same protein-readout collapse (around 0.08 at the 25% "
  "holdout), indicating the loss is intrinsic to RNA-to-protein transfer rather than an under-training "
  "artefact."),
 # Frangieh seed-range qualifier (table cells, x2)
 ("seed-robust {0,1,2}", "seed range over {0,1,2} (not a cluster CI)"),
 # CellOT per-fold adapted/baseline note (Table 6 section footnote)
 ("These are additional metric checks, not distributional comparators.",
  "These are additional metric checks, not distributional comparators. On Soskic leave-one-donor-out, "
  "CellOT and the per-donor primary baseline are both refit within each held-out fold from training "
  "cells only, whereas scGen is native (no per-fold retraining) and on the other axes the primary "
  "baseline is pooled from the training split; the per-fold CellOT and its baseline use the same "
  "training-only response genes and PCA basis, so the comparison is leak-free and symmetric."),
]


def iter_paras(doc):
    for p in doc.paragraphs:
        yield p
    for el in doc.element.body.iterchildren():
        if el.tag.split("}")[-1] == "tbl":
            for row in Table(el, doc).rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def apply_pair(doc, find, repl):
    hits = 0
    for p in iter_paras(doc):
        joined = "".join(r.text for r in p.runs)
        if find in joined:
            new = joined.replace(find, repl)
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            hits += 1
    return hits


def main():
    bak = DOCX.with_name(DOCX.stem + ".pre-disclosures.docx")
    if not bak.exists():
        shutil.copy2(DOCX, bak)
    doc = Document(str(DOCX))
    misses = []
    for find, repl in PAIRS:
        h = apply_pair(doc, find, repl)
        print(f"  {h:>2}x  {find[:48]!r}")
        if h == 0:
            misses.append(find[:48])
    if misses:
        print(f"!! NOT MATCHED: {misses}", file=sys.stderr)
        sys.exit(1)
    doc.save(str(DOCX))
    print(f"saved {DOCX.name}; backup {bak.name}")


if __name__ == "__main__":
    main()
