#!/usr/bin/env python3
"""Insert the new Results content (3 design tables + bridge/coverage/roster/Limitations paragraphs) into
both .docx, and apply this round's renumber/caveat edits. Incremental (does NOT rebuild the section) so
the already-styled prose, figures, and Table 6 are preserved.

New tables are built in the manuscript table style (fixed layout, borderless 'Normal Table', Times New
Roman, bold header), matching Table 6 / the Methods tables. Content matches results_section.md.
Run BEFORE update_docx_captions.py + style_results_in_manuscript.py + fix_table3.py (which then refresh
the figure captions, set TNR sizes across the region, and re-fix Table 6 geometry). Idempotent-guarded:
skips a table insert if its header already exists.
"""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Twips
from docx.table import Table
from docx.text.paragraph import Paragraph

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
WD = Path(__file__).resolve().parents[1]
FONT = "Times New Roman"
BODY_PT, SUBHEAD_PT, CAP_PT, TBL_PT = 10.0, 12.0, 9.0, 9.0


def qn(t): return W + t


# ---------- content (matches results_section.md) ----------
BRIDGE = ("These five task clusters map onto the four generalization axes through five benchmark "
    "components (Table 3). Cytokine-response prediction (Kang) supplies the cell-context split and a "
    "random-split donor control; donor-aware activation dynamics is evaluated directly on Soskic CD4 "
    "activation with leave-one-donor-out; gene-intervention prediction (the primary-human-T CRISPR "
    "collection) supplies the perturbation axis; complex-context perturbation (Frangieh) supplies the "
    "modality axis; and small-molecule perturbation (OP3) supplies a second cell-context split together "
    "with the unseen-compound perturbation split. The previous Soskic exclusion reason does not hold for "
    "the processed 0h-versus-16h data: paired control-perturbed donor structure and matched simple "
    "baselines are present. Every axis therefore has at least one evaluated dataset, and no benchmark "
    "component is left unmapped.")
T3_CAP = ("**Table 3. Mapping of the five review task clusters onto the four generalization axes and the "
    "five benchmark components.**")
T3 = [["Task cluster (review framework)", "Benchmark component", "Generalization axis (evaluation split)"],
    ["Cytokine-response prediction", "Kang IFN-β PBMC", "Cell-context (leave-one-lineage-out) + donor (leave-one-donor-out)"],
    ["Donor-aware activation dynamics", "Soskic CD4 activation (106 paired donors)", "Donor (leave-one-donor-out)"],
    ["Gene-intervention in primary immune cells", "primary-human-T CRISPR collection", "Perturbation (leave-one-gene-out)"],
    ["Complex-context perturbation", "Frangieh Perturb-CITE-seq", "Modality (RNA vs CITE-seq surface protein)"],
    ["Small-molecule perturbation", "OP3 PBMC compound atlas", "Cell-context (leave-one-lineage-out) + perturbation (unseen compound)"]]

METRIC_PARA = ("The four immune-aware metrics are not all defined on every component (Table 4). "
    "Response-direction fidelity (Pearson-Δ) and distributional fidelity (E-distance) are computed on all "
    "five datasets. Immune-program fidelity (AUCell-Δ) is computed wherever a curated transcriptional immune "
    "gene set is defined — the type-I-interferon set for Kang, T-cell activation / IL2-STAT5 / type-I and "
    "type-II IFN modules for Soskic, three response programs (ISG, NF-κB, effector) for OP3, and five "
    "T-cell programs for the primary-T CRISPR collection — but not for Frangieh, whose held readout is the "
    "CITE-seq surface-protein panel rather than a transcriptional program. Split-stratified generalization "
    "robustness applies to every endpoint by construction. The Figure 2 schematic therefore depicts the "
    "full metric menu, while Table 4 records where each metric is actually instantiated.")
T4_CAP = "**Table 4. Metric coverage: which immune-aware metric is computed on each benchmark component.**"
T4 = [["Benchmark component", "Pearson-Δ", "E-distance", "AUCell-Δ (immune program)", "Robustness (split-stratified)"],
    ["Kang IFN-β PBMC", "✓", "✓", "✓ (type-I-IFN)", "✓"],
    ["Soskic CD4 activation", "✓", "✓", "✓ (T-cell activation; IL2/STAT5; type-I/II IFN)", "✓"],
    ["OP3 compound atlas", "✓", "✓", "✓ (ISG; NF-κB; effector)", "✓"],
    ["primary-T CRISPR collection", "✓", "✓", "✓ (5 T-cell programs)", "✓"],
    ["Frangieh Perturb-CITE-seq", "✓", "✓", "— (surface-protein readout; no transcriptional gene set)", "✓"]]

ROSTER_PARA = ("The benchmark is a deliberate subset of the methods surveyed in the preceding section, not "
    "an exhaustive bake-off. Of the roughly thirty perturbation-prediction methods catalogued there, eight "
    "conditioned models spanning six families were executed here (Table 5), each chosen for having a "
    "maintained implementation and meeting a given split's data requirements. A model is reported as native "
    "where its published task matches the split, adapted where it was re-fit to a split outside its original "
    "design, and as a floor where only its non-conditioned reduction applies; CINEMA-OT is carried as a "
    "perturbation-agnostic distributional comparator rather than a conditioned predictor. A surveyed method "
    "was not run when no maintained implementation was available, when a split's inputs were undefined for "
    "that dataset (for example, a chemical structure for an unseen compound, or a gene-perturbation graph "
    "for a graph model), or when the method does not apply to single-cell perturbation transfer. These "
    "omissions reflect method–task compatibility rather than model quality, and recording them makes the "
    "funnel from survey to benchmark explicit and intentional.")
T5_CAP = ("**Table 5. Conditioned models executed in the benchmark: family, datasets and splits run, and "
    "applicability.**")
T5 = [["Model", "Family", "Run on (dataset / split)", "Applicability"],
    ["scGen", "latent", "Kang, Soskic, Frangieh; CRISPR, OP3", "native (Kang, Soskic, Frangieh) / adapted (CRISPR, OP3)"],
    ["CPA", "latent", "Kang, OP3; CRISPR", "native (Kang, OP3) / adapted (CRISPR)"],
    ["GEARS", "graph", "CRISPR leave-one-gene-out", "native"],
    ["AttentionPert", "graph", "CRISPR leave-one-gene-out", "native"],
    ["scGPT", "foundation", "CRISPR leave-one-gene-out", "native"],
    ["STATE", "hybrid", "CRISPR; OP3", "native (CRISPR) / adapted (OP3)"],
    ["FP-ridge", "chemistry", "OP3 leave-one-lineage-out + unseen compound", "native"],
    ["linear-shift-KOemb", "shift", "Frangieh unseen knockout", "native"],
    ["CINEMA-OT", "optimal-transport", "CRISPR, OP3", "distributional comparator (not a conditioned win)"]]

LIM_PARA = ("Several limitations bound these conclusions. First, statistical power is modest: the resampling "
    "unit is the biological replicate — 6 immune lineages, 5 datasets, 106 Soskic donors, or 8 Kang donors "
    "for the random-split control — so the axis-level endpoints are consistent-direction effects rather "
    "than high-powered tests. Second, the evaluated panel is a deliberate funnel from the surveyed landscape "
    "(Tables 3–5): datasets were admitted only where paired control-perturbed structure and matched "
    "baselines existed; Soskic now satisfies this criterion for the 0h-versus-16h donor axis, whereas "
    "temporal saturation and several adjacent resources remain outside the present benchmark. Third, model "
    "coverage is partial — eight conditioned models of the roughly thirty surveyed were executed (Table 5), "
    "so the absence of a published method here is not evidence against it. Fourth, the single positive axis "
    "(seen-compound cell-context transfer on OP3) is also the comparatively easier task — interpolating an "
    "already-seen perturbation across lineages — so it should not be read as general evidence that "
    "conditioning transfers.")

REPLACEMENTS = [
    ("reported as exploratory only.",
     "reported as exploratory only. Kang is retained as the random-split optimism control; the dedicated "
     "donor-axis evidence comes from Soskic CD4 activation leave-one-donor-out over 106 processed paired "
     "donors."),
    ("for unseen-gene, unseen-compound, modality, or donor generalization in these splits.",
     "for unseen-gene, unseen-compound, or modality generalization in these splits, nor for donor "
     "generalization in the Soskic leave-one-donor-out analysis."),
]


# ---------- helpers ----------
def _runs_from_md(paragraph, text):
    for i, seg in enumerate(text.split("**")):
        if seg == "":
            continue
        r = paragraph.add_run(seg)
        r.bold = (i % 2 == 1)
        r.font.name = FONT


def make_para(doc, text, size):
    p = doc.add_paragraph()
    _runs_from_md(p, text)
    for r in p.runs:
        r.font.size = Pt(size)
    return p._p


def make_table(doc, rows, table_w, fracs, font_pt=TBL_PT):
    ncol = len(rows[0])
    widths = [max(400, round(table_w * f)) for f in fracs]
    widths[-1] += table_w - sum(widths)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Normal Table"
    tblPr = t._tbl.find(qn("tblPr"))
    tw = tblPr.find(qn("tblW"));  tw.set(qn("w"), str(table_w)); tw.set(qn("type"), "dxa")
    from lxml import etree
    lay = tblPr.find(qn("tblLayout")) or etree.SubElement(tblPr, qn("tblLayout"))
    lay.set(qn("type"), "fixed")
    grid = t._tbl.find(qn("tblGrid"))
    for c, w in zip(grid.findall(qn("gridCol")), widths):
        c.set(qn("w"), str(w))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.width = Twips(widths[ci])
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(font_pt)
            if ri == 0:
                run.font.bold = True
    return t._tbl


def find_para(doc, pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p
    return None


def replace_text(doc, old, new):
    n = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new); n += 1
    for el in doc.element.body.iterchildren():
        if el.tag == qn("tbl"):
            for row in Table(el, doc).rows:
                for cell in row.cells:
                    for pp in cell.paragraphs:
                        for r in pp.runs:
                            if old in r.text:
                                r.text = r.text.replace(old, new); n += 1
    return n


def table_exists(doc, header0):
    for el in doc.element.body.iterchildren():
        if el.tag == qn("tbl") and Table(el, doc).rows[0].cells[0].text.strip() == header0:
            return True
    return False


def process(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    sec = doc.sections[0]
    table_w = int((sec.page_width - sec.left_margin - int(0.28 * 914400)) / 635)

    if table_exists(doc, "Task cluster (review framework)"):
        print(f"  {docx_path.name}: design tables already present; skipping insert")
        return

    # 1) renumber the decision table (Table 3 -> Table 6) + apply caveat/Synthesis edits BEFORE inserting
    replace_text(doc, "Table 3", "Table 6")
    for old, new in REPLACEMENTS:
        replace_text(doc, old, new)

    # 2) insert content. helper: chain elements after a cursor element.
    def chain_after(cursor, elements):
        for el in elements:
            cursor.addnext(el); cursor = el
        return cursor

    fr3 = [0.32, 0.27, 0.41]
    fr4 = [0.21, 0.09, 0.10, 0.37, 0.23]
    fr5 = [0.15, 0.13, 0.39, 0.33]

    # A1: after the Figure 2 caption -> bridge + Table 3
    a1 = find_para(doc, lambda t: t.startswith("Figure 2."))
    chain_after(a1._p, [make_para(doc, BRIDGE, BODY_PT),
                        make_para(doc, T3_CAP, CAP_PT),
                        make_table(doc, T3, table_w, fr3)])

    # A2: after the "Throughout, hypothesis tests" paragraph -> metric+T4 + roster+T5
    a2 = find_para(doc, lambda t: t.startswith("Every model was scored against four simple baselines"))
    chain_after(a2._p, [make_para(doc, METRIC_PARA, BODY_PT),
                        make_para(doc, T4_CAP, CAP_PT),
                        make_table(doc, T4, table_w, fr4),
                        make_para(doc, ROSTER_PARA, BODY_PT),
                        make_para(doc, T5_CAP, CAP_PT),
                        make_table(doc, T5, table_w, fr5)])

    # A3: before the Synthesis heading -> Limitations heading + paragraph
    syn = find_para(doc, lambda t: t == "Synthesis")
    lim_h = make_para(doc, "Limitations", SUBHEAD_PT)
    for r in Paragraph(lim_h, doc).runs:
        r.font.bold = True
    lim_p = make_para(doc, LIM_PARA, BODY_PT)
    syn._p.addprevious(lim_h)
    syn._p.addprevious(lim_p)

    doc.save(str(docx_path))
    print(f"  {docx_path.name}: inserted 3 tables + bridge/metric/roster/Limitations; renumbered Table 3->6")


if __name__ == "__main__":
    for p in [WD / "results" / "_paper" / "results_section.docx",
              WD.parent / "Toward Immune Virtual Cells (benchmark Results inserted).docx"]:
        if not p.exists():
            sys.exit(f"FAIL: docx not found: {p}")
        print(f"adding Results content to {p}")
        process(p)
    print("done.")
