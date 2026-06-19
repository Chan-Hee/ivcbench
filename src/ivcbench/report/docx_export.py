"""Export an auto-generated draft subsection to Word (.docx), with the figure embedded.

Produces a paste-ready subsection for "Toward Immune Virtual Cells": heading, prose (with bold
lead-ins preserved), the results table, and the figure with a caption. Light Markdown parsing tuned
to what report.draft emits.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

_FIG_NO = {"C1": "Figure 3", "C2": "Figure 4", "C3": "Figure 5", "C4": "Figure 6", "C5": "Figure 7"}
_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _add_runs(paragraph, text: str) -> None:
    """Render inline **bold** and *italic* Markdown as styled runs."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            paragraph.add_run(tok[1:-1]).italic = True
        else:
            paragraph.add_run(tok)


def build_docx(cluster: str, draft_md: str, figure_png: str | Path | None, out_path: str | Path) -> Path:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    table_buffer: list[str] = []
    para_buffer: list[str] = []

    def flush_table():
        if not table_buffer:
            return
        rows = [r for r in table_buffer if not set(r.replace("|", "").strip()) <= {"-", " "}]
        cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
        table_buffer.clear()
        if not cells:
            return
        t = doc.add_table(rows=len(cells), cols=len(cells[0]))
        t.style = "Light Grid Accent 1"
        for ri, row in enumerate(cells):
            for ci, val in enumerate(row):
                run = t.cell(ri, ci).paragraphs[0].add_run(val)
                run.font.size = Pt(8.5)
                if ri == 0:
                    run.bold = True

    def flush_para():
        if para_buffer:  # soft newlines within a markdown paragraph join with a space
            _add_runs(doc.add_paragraph(), " ".join(para_buffer))
            para_buffer.clear()

    for line in draft_md.splitlines():
        s = line.rstrip()
        if s.lstrip().startswith("|"):
            flush_para()
            table_buffer.append(s)
            continue
        flush_table()
        if not s.strip():
            flush_para()
        elif s.startswith("> "):  # synthetic-data / provenance banner -> grey italic note
            flush_para()
            r = doc.add_paragraph().add_run(s[2:].replace("**", ""))
            r.italic = True
            r.font.color.rgb = RGBColor(0x99, 0x55, 0x00)
            r.font.size = Pt(9)
        elif s.startswith("#### "):
            flush_para(); doc.add_heading(s[5:], level=3)
        elif s.startswith("### "):
            flush_para(); doc.add_heading(s[4:], level=2)
        elif s.startswith("## "):
            flush_para(); doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            flush_para(); doc.add_heading(s[2:], level=1)
        elif s.lstrip().startswith(("- ", "* ")):  # bullet list item
            flush_para()
            _add_runs(doc.add_paragraph(style="List Bullet"), s.lstrip()[2:])
        else:
            para_buffer.append(s.strip())
    flush_para()
    flush_table()

    figure_png = Path(figure_png) if figure_png else None
    if figure_png and figure_png.exists():
        doc.add_picture(str(figure_png), width=Inches(6.5))
        cap = doc.add_paragraph()
        cr = cap.add_run(f"{_FIG_NO.get(cluster, 'Figure')}. {cluster} benchmark — four evaluation "
                         "axes under leak-proof, applicability-gated splits.")
        cr.italic = True
        cr.font.size = Pt(9)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
